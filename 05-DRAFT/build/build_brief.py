#!/usr/bin/env python3
"""
Render the executive summary, PDF and DOCX, three languages.

Deliberately NOT the paper's layout. This is a policy brief: wider measure,
larger type, numbered finding blocks, a pull quote, an accent rule down the left
of callouts. Built to be forwarded and read in five minutes.

    BRIEF_LANG=en|it|fr python3 build_brief.py
"""
import os, re

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                Spacer, Image, Table, TableStyle, KeepTogether,
                                HRFlowable)
from PIL import Image as PILImage

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import brief_content as BC
from build_pdf import BASE, BOLD, ITAL, rt

LANG = os.environ.get("BRIEF_LANG", "en").lower()
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(HERE, "..", "..", "06-SUBMISSION"))
os.makedirs(OUT, exist_ok=True)

META, BODY = BC.build(LANG)

INK = colors.HexColor("#141414")
GREY = colors.HexColor("#666666")
LGREY = colors.HexColor("#9e9e9e")
ACCENT = colors.HexColor("#8c2d04")
WASH = colors.HexColor("#f6f4f1")
RULE = colors.HexColor("#d8d3cc")

PW, PH = A4
LM = RM = 24 * mm
TM, BM = 20 * mm, 18 * mm
CW = PW - LM - RM


def S(name, **kw):
    b = dict(name=name, fontName=BASE, fontSize=10.2, leading=14.4, textColor=INK,
             alignment=TA_JUSTIFY, spaceAfter=0)
    b.update(kw)
    return ParagraphStyle(**b)


BI = "TNR-BI" if BASE == "TNR" else "Times-BoldItalic"
st_kick = S("k", fontName=BOLD, fontSize=8.2, leading=10, textColor=ACCENT, alignment=TA_LEFT)
st_title = S("t", fontName=BOLD, fontSize=22, leading=25, alignment=TA_LEFT, spaceAfter=4)
st_sub = S("s", fontName=ITAL, fontSize=11.6, leading=15, textColor=GREY, alignment=TA_LEFT)
st_by = S("b", fontSize=8.8, leading=11.5, textColor=GREY, alignment=TA_LEFT)
st_lead = S("l", fontSize=11.6, leading=16.2, spaceAfter=9)
st_h = S("h", fontName=BOLD, fontSize=13, leading=16, alignment=TA_LEFT, spaceAfter=5)
st_p = S("p", spaceAfter=7)
st_pull = S("q", fontName=BI, fontSize=13, leading=17.5, textColor=ACCENT, alignment=TA_LEFT)
st_fno = S("fn", fontName=BOLD, fontSize=21, leading=22, textColor=ACCENT, alignment=TA_LEFT)
st_ft = S("ft", fontName=BOLD, fontSize=11.4, leading=14.4, alignment=TA_LEFT, spaceAfter=3)
st_bul = S("bl", leftIndent=12, bulletIndent=2, spaceAfter=5)
st_cap = S("c", fontSize=8.2, leading=10.8, textColor=GREY, alignment=TA_LEFT)
st_boxt = S("bt", fontName=BOLD, fontSize=10.2, leading=13, alignment=TA_LEFT, spaceAfter=3)
st_box = S("bx", fontSize=9.4, leading=13)
st_note = S("n", fontSize=7.8, leading=10, textColor=LGREY, alignment=TA_LEFT)


def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont(BASE, 7.6); canvas.setFillColor(LGREY)
    canvas.drawString(LM, BM - 8 * mm, META["running"])
    canvas.drawRightString(PW - RM, BM - 8 * mm, str(canvas.getPageNumber()))
    canvas.setStrokeColor(RULE); canvas.setLineWidth(.5)
    canvas.line(LM, BM - 5.5 * mm, PW - RM, BM - 5.5 * mm)
    canvas.restoreState()


def fig(path, cap):
    if not os.path.exists(path):
        return [Paragraph(rt(f"*[missing {os.path.basename(path)}]*"), st_note)]
    w, h = PILImage.open(path).size
    tw = CW
    img = Image(path, width=tw, height=tw * h / w)
    return [Spacer(1, 4),
            KeepTogether([img, Spacer(1, 3), Paragraph(rt(cap), st_cap)]),
            Spacer(1, 9)]


def finding(no, title, body):
    left = Table([[Paragraph(no, st_fno)]], colWidths=[13 * mm])
    left.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"),
                              ("LEFTPADDING", (0, 0), (-1, -1), 0),
                              ("TOPPADDING", (0, 0), (-1, -1), 0)]))
    right = Table([[Paragraph(rt(title), st_ft)], [Paragraph(rt(body), st_p)]],
                  colWidths=[CW - 13 * mm])
    right.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0),
                               ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                               ("TOPPADDING", (0, 0), (-1, -1), 0),
                               ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
    t = Table([[left, right]], colWidths=[13 * mm, CW - 13 * mm])
    t.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"),
                           ("LEFTPADDING", (0, 0), (-1, -1), 0),
                           ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                           ("TOPPADDING", (0, 0), (-1, -1), 2),
                           ("BOTTOMPADDING", (0, 0), (-1, -1), 2)]))
    return [Spacer(1, 3), t, Spacer(1, 2)]


def build_pdf(path):
    doc = BaseDocTemplate(path, pagesize=A4, leftMargin=LM, rightMargin=RM,
                          topMargin=TM, bottomMargin=BM, title=META["title"],
                          author=META["author"], subject=META["subtitle"])
    fr = Frame(LM, BM, CW, PH - TM - BM, id="f", leftPadding=0, rightPadding=0,
               topPadding=0, bottomPadding=0)
    doc.addPageTemplates([PageTemplate(id="p", frames=[fr], onPage=_footer)])

    F = [Paragraph(META["kicker"], st_kick), Spacer(1, 3),
         Paragraph(rt(META["title"]), st_title),
         Paragraph(rt(META["subtitle"]), st_sub), Spacer(1, 7),
         Paragraph(f'{META["author"]} &nbsp;·&nbsp; {META["affil"]} '
                   f'&nbsp;·&nbsp; {META["date"]}', st_by),
         Spacer(1, 5),
         HRFlowable(width="100%", thickness=1.4, color=ACCENT, spaceAfter=10)]

    for blk in BODY:
        k = blk[0]
        if k == "lead":
            F += [Paragraph(rt(blk[1]), st_lead)]
        elif k == "h":
            F += [Spacer(1, 5), Paragraph(rt(blk[1]), st_h)]
        elif k == "p":
            F += [Paragraph(rt(blk[1]), st_p)]
        elif k == "rule":
            F += [Spacer(1, 3), HRFlowable(width="100%", thickness=.6, color=RULE,
                                           spaceAfter=8)]
        elif k == "pull":
            t = Table([[Paragraph(rt(blk[1]), st_pull)]], colWidths=[CW])
            t.setStyle(TableStyle([("LINEBEFORE", (0, 0), (0, 0), 2.4, ACCENT),
                                   ("LEFTPADDING", (0, 0), (-1, -1), 11),
                                   ("TOPPADDING", (0, 0), (-1, -1), 5),
                                   ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
            F += [Spacer(1, 3), t, Spacer(1, 9)]
        elif k == "finding":
            F += finding(blk[1], blk[2], blk[3])
        elif k == "bullets":
            for it in blk[1]:
                F += [Paragraph(rt(it), st_bul, bulletText="•")]
            F += [Spacer(1, 4)]
        elif k == "fig":
            F += fig(blk[1], blk[2])
        elif k == "box":
            inner = Table([[Paragraph(rt(blk[1]), st_boxt)],
                           [Paragraph(rt(blk[2]), st_box)]], colWidths=[CW - 18])
            inner.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0),
                                       ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                                       ("TOPPADDING", (0, 0), (-1, -1), 0),
                                       ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
            t = Table([[inner]], colWidths=[CW])
            t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), WASH),
                                   ("LEFTPADDING", (0, 0), (-1, -1), 11),
                                   ("RIGHTPADDING", (0, 0), (-1, -1), 11),
                                   ("TOPPADDING", (0, 0), (-1, -1), 9),
                                   ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                                   ("LINEBEFORE", (0, 0), (0, 0), 2.4, ACCENT)]))
            F += [Spacer(1, 5), KeepTogether(t)]

    F += [Spacer(1, 8), Paragraph(rt(META["note"]), st_note)]
    doc.build(F)
    return path


# ── DOCX ─────────────────────────────────────────────────────────────────────
D_INK, D_GREY, D_ACC = RGBColor(0x14, 0x14, 0x14), RGBColor(0x66, 0x66, 0x66), RGBColor(0x8C, 0x2D, 0x04)
TOK = re.compile(r"(\*\*.+?\*\*|(?<![\w*])\*[^*\n]+?\*(?![\w*]))")


def rich(par, text, size=10.2, color=D_INK, italic=False, bold=False):
    for piece in TOK.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            r = par.add_run(piece[1:-1]); r.italic = True
        else:
            r = par.add_run(piece); r.italic = italic; r.bold = bold
        r.font.size = Pt(size); r.font.color.rgb = color
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def dpara(doc, text="", size=10.2, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=7,
          color=D_INK, italic=False, bold=False, style=None):
    p = doc.add_paragraph(style=style); p.alignment = align
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.16
    if text:
        rich(p, text, size, color, italic, bold)
    return p


def dshade(cell, hexc):
    e = OxmlElement("w:shd"); e.set(qn("w:val"), "clear"); e.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(e)


def dborder(par, pos, sz, color):
    pPr = par._p.get_or_add_pPr()
    bd = pPr.find(qn("w:pBdr"))
    if bd is None:
        bd = OxmlElement("w:pBdr"); pPr.append(bd)
    e = OxmlElement(f"w:{pos}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "4"); e.set(qn("w:color"), color)
    bd.append(e)


def build_docx(path):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.4)
    s.top_margin, s.bottom_margin = Cm(2.0), Cm(1.8)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"; st.font.size = Pt(10.2)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    ft = s.footer.paragraphs[0]; ft.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = ft.add_run(META["running"]); r.font.size = Pt(7.6)
    r.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E); r.font.name = "Times New Roman"

    dpara(doc, META["kicker"], 8.2, WD_ALIGN_PARAGRAPH.LEFT, 3, D_ACC, bold=True)
    t = dpara(doc, META["title"], 22, WD_ALIGN_PARAGRAPH.LEFT, 4, bold=True)
    dpara(doc, META["subtitle"], 11.6, WD_ALIGN_PARAGRAPH.LEFT, 7, D_GREY, italic=True)
    by = dpara(doc, f'{META["author"]}  ·  {META["affil"]}  ·  {META["date"]}',
               8.8, WD_ALIGN_PARAGRAPH.LEFT, 8, D_GREY)
    dborder(by, "bottom", 16, "8C2D04")

    for blk in BODY:
        k = blk[0]
        if k == "lead":
            dpara(doc, blk[1], 11.6, after=9)
        elif k == "h":
            h = dpara(doc, blk[1], 13, WD_ALIGN_PARAGRAPH.LEFT, 5, bold=True)
            h.paragraph_format.space_before = Pt(9)
            h.paragraph_format.keep_with_next = True
        elif k == "p":
            dpara(doc, blk[1])
        elif k == "rule":
            rp = dpara(doc, "", after=6); dborder(rp, "bottom", 4, "D8D3CC")
        elif k == "pull":
            q = dpara(doc, blk[1], 13, WD_ALIGN_PARAGRAPH.LEFT, 9, D_ACC, italic=True)
            q.paragraph_format.left_indent = Cm(0.4)
            dborder(q, "left", 18, "8C2D04")
        elif k == "finding":
            fp = dpara(doc, "", 21, WD_ALIGN_PARAGRAPH.LEFT, 1, D_ACC, bold=True)
            rich(fp, f"{blk[1]}   ", 21, D_ACC, bold=True)
            rich(fp, blk[2], 11.4, D_INK, bold=True)
            fp.paragraph_format.keep_with_next = True
            b = dpara(doc, blk[3], after=8)
            b.paragraph_format.left_indent = Cm(0.0)
        elif k == "bullets":
            for it in blk[1]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                bp.paragraph_format.space_after = Pt(5)
                bp.paragraph_format.left_indent = Cm(0.75)
                rich(bp, it)
        elif k == "fig":
            if os.path.exists(blk[1]):
                ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_before = Pt(5)
                ip.paragraph_format.space_after = Pt(3)
                ip.paragraph_format.keep_with_next = True
                w, h = PILImage.open(blk[1]).size
                cm = 16.2
                ip.add_run().add_picture(blk[1], width=Cm(cm), height=Cm(cm * h / w))
                dpara(doc, blk[2], 8.2, WD_ALIGN_PARAGRAPH.LEFT, 9, D_GREY)
        elif k == "box":
            tb = doc.add_table(rows=1, cols=1); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            c = tb.cell(0, 0); dshade(c, "F6F4F1")
            p1 = c.paragraphs[0]; p1.paragraph_format.space_after = Pt(3)
            rich(p1, blk[1], 10.2, D_INK, bold=True)
            p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            rich(p2, blk[2], 9.4)

    dpara(doc, META["note"], 7.8, WD_ALIGN_PARAGRAPH.LEFT, 0,
          RGBColor(0x9E, 0x9E, 0x9E))
    doc.save(path)
    return path


STEM = {"en": "Escoda-2026-Half-Life-of-Compute-EXECUTIVE-SUMMARY",
        "it": "Escoda-2026-Emivita-del-Calcolo-SINTESI-ESECUTIVA-IT",
        "fr": "Escoda-2026-Demi-Vie-du-Calcul-SYNTHESE-FR"}

if __name__ == "__main__":
    stem = STEM[LANG]
    p = build_pdf(os.path.join(OUT, f"{stem}.pdf"))
    print(f"BRIEF PDF  [{LANG}] → {os.path.basename(p)} ({os.path.getsize(p)/1024:.0f} KB)")
    p = build_docx(os.path.join(OUT, f"{stem}.docx"))
    print(f"BRIEF DOCX [{LANG}] → {os.path.basename(p)} ({os.path.getsize(p)/1024:.0f} KB)")
