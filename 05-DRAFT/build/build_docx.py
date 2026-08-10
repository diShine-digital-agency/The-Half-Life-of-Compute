#!/usr/bin/env python3
"""
Render the manuscript to an editable .docx (python-docx).

Same content module as the PDF, so the two cannot diverge. Styling mirrors the
PDF: Times New Roman, A4, justified body, booktabs-style table rules, numbered
figure/table captions. Equations are embedded as images (Word has no LaTeX).
"""
import os, re

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

import importlib
LANG = os.environ.get("PAPER_LANG", "en").lower()
C = importlib.import_module("paper_content" if LANG == "en"
                            else f"paper_content_{LANG}")
from build_pdf import eq_png, eq_lines, GREEK

L_ABS = C.META.get("abstract_head", "ABSTRACT")
L_KW = C.META.get("kw_label", "Keywords")
L_FIG = C.META.get("fig_label", "Figure")
L_TAB = C.META.get("tab_label", "Table")
L_RUN = C.META.get("running", "Escoda · The Half-Life of Compute")

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(HERE, "..", "..", "06-SUBMISSION"))
os.makedirs(OUT, exist_ok=True)

INK = RGBColor(0x11, 0x11, 0x11)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
LGREY = RGBColor(0x9A, 0x9A, 0x9A)
ACCENT = RGBColor(0x8C, 0x2D, 0x04)
BODY_PT, CONTENT_W_CM = 10.0, 15.6


# ── inline markup → runs ─────────────────────────────────────────────────────
def math_to_text(s):
    for k, v in GREEK.items():
        s = s.replace(k, v)
    s = re.sub(r"\^\{([^{}]*)\}", r"^\1", s)
    s = re.sub(r"_\{([^{}]*)\}", r"_\1", s)
    return s.replace("{", "").replace("}", "").replace("\\", "")


TOK = re.compile(r"(\$[^$]+\$|\*\*.+?\*\*|(?<![\w*])\*[^*\n]+?\*(?![\w*]))")


def add_rich(par, text, size=BODY_PT, color=INK, base_italic=False):
    """Split on $math$, **bold**, *italic* and emit styled runs."""
    for piece in TOK.split(text):
        if not piece:
            continue
        if piece.startswith("$") and piece.endswith("$"):
            r = par.add_run(math_to_text(piece[1:-1])); r.italic = True
        elif piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            r = par.add_run(piece[1:-1]); r.italic = True
        else:
            r = par.add_run(piece); r.italic = base_italic
        r.font.size = Pt(size); r.font.color.rgb = color
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return par


def para(doc, text="", size=BODY_PT, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=6, color=INK, italic=False, indent_left=0, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(0)
    pf.line_spacing = 1.13
    if indent_left:
        pf.left_indent = Cm(indent_left)
    if text:
        add_rich(p, text, size, color, italic)
    return p


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def hrule(par, pos="bottom", sz=6, color="333333"):
    pPr = par._p.get_or_add_pPr()
    bd = pPr.find(qn("w:pBdr"))
    if bd is None:
        bd = OxmlElement("w:pBdr"); pPr.append(bd)
    e = OxmlElement(f"w:{pos}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "2"); e.set(qn("w:color"), color)
    bd.append(e)


def cell_border(cell, pos, sz=6, color="333333"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = OxmlElement("w:tcBorders"); tcPr.append(b)
    e = OxmlElement(f"w:{pos}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz)); e.set(qn("w:color"), color)
    b.append(e)


def img_w_h(path, target_cm):
    w, h = PILImage.open(path).size
    return Cm(target_cm), Cm(target_cm * h / w)


# ── document ─────────────────────────────────────────────────────────────────
def build(path):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.7)
    s.top_margin, s.bottom_margin = Cm(2.2), Cm(2.0)

    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"; st.font.size = Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    M = C.META
    # running header
    hdr = s.header.paragraphs[0]
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hdr.add_run(f"{L_RUN}    ·    {M['date']}")
    hr.font.size = Pt(7.6); hr.font.color.rgb = LGREY; hr.font.name = "Times New Roman"

    # ── front matter ─────────────────────────────────────────────────────────
    t = para(doc, M["title"], size=17, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=9)
    for r in t.runs:
        r.bold = True; r.font.size = Pt(17)

    para(doc, M["authors"], size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, f"{M['affiliation']}  ·  {M['email']}", size=9.2,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, color=GREY, italic=True)
    para(doc, f"ORCID {M['orcid']}", size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=1, color=GREY, italic=True)
    para(doc, M['date'], size=8.8,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, color=GREY)

    dp = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    for txt, bold in [("DOI  ", True), (M['doi'], False)]:
        r = dp.add_run(txt); r.bold = bold; r.font.size = Pt(8.6)
        r.font.color.rgb = GREY; r.font.name = "Times New Roman"
    hrule(dp, "top", 4, "CCCCCC"); hrule(dp, "bottom", 4, "CCCCCC")

    # abstract in a shaded single-cell table
    at = doc.add_table(rows=1, cols=1); at.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = at.cell(0, 0); shade(cell, "F7F7F5")
    for pos in ("top", "bottom"):
        cell_border(cell, pos, 8, "333333")
    ap = cell.paragraphs[0]; ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ap.add_run(L_ABS); r.bold = True; r.font.size = Pt(9.4)
    r.font.name = "Times New Roman"
    ap.paragraph_format.space_after = Pt(4)
    bp = cell.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bp.paragraph_format.line_spacing = 1.1
    add_rich(bp, C.ABSTRACT, 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    kp = para(doc, "", space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = kp.add_run(f"{L_KW}  "); r.bold = True; r.font.size = Pt(8.6); r.font.color.rgb = GREY
    r.font.name = "Times New Roman"
    add_rich(kp, M["keywords"], 8.6, GREY)
    jp = para(doc, "", space_after=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    for txt, bold in [("JEL  ", True), (f"{M['jel']}    ·    ", False),
                      ("ACM CCS  ", True), (M["acm"], False)]:
        r = jp.add_run(txt); r.bold = bold; r.font.size = Pt(8.6)
        r.font.color.rgb = GREY; r.font.name = "Times New Roman"

    # ── body ─────────────────────────────────────────────────────────────────
    for blk in C.ALL:
        k = blk[0]

        if k == "h1":
            h = para(doc, "", space_after=5, align=WD_ALIGN_PARAGRAPH.LEFT)
            h.paragraph_format.space_before = Pt(11); h.paragraph_format.keep_with_next = True
            r = h.add_run(blk[1]); r.bold = True; r.font.size = Pt(11.8)
            r.font.name = "Times New Roman"
            hrule(h, "bottom", 4, "DDDDDD")

        elif k == "h2":
            h = para(doc, "", space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
            h.paragraph_format.space_before = Pt(7); h.paragraph_format.keep_with_next = True
            r = h.add_run(blk[1]); r.bold = True; r.italic = True; r.font.size = Pt(10.4)
            r.font.name = "Times New Roman"

        elif k == "p":
            para(doc, blk[1])

        elif k == "quote":
            q = para(doc, blk[1], italic=True, indent_left=0.5, space_after=8, color=INK)
            hrule(q, "left", 18, "8C2D04")

        elif k in ("bullets", "numbers"):
            for i, it in enumerate(blk[1], 1):
                bp = doc.add_paragraph(
                    style="List Bullet" if k == "bullets" else "List Number")
                bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                bp.paragraph_format.space_after = Pt(3.4)
                bp.paragraph_format.left_indent = Cm(0.75)
                add_rich(bp, it)

        elif k == "eq":
            for i, ln in enumerate(eq_lines(blk[1])):
                png = eq_png(ln, f"{blk[2].replace('.', '_')}_{i}")
                ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ep.paragraph_format.space_after = Pt(2)
                w, h = PILImage.open(png).size
                cm = min(w / 190.0, CONTENT_W_CM * 0.7)
                ep.add_run().add_picture(png, width=Cm(cm), height=Cm(cm * h / w))
                if i == len(eq_lines(blk[1])) - 1:
                    tr = ep.add_run(f"      ({blk[2]})")
                    tr.font.size = Pt(9); tr.font.name = "Times New Roman"
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif k == "fig":
            path_, cap, lab = blk[1], blk[2], blk[3]
            if os.path.exists(path_):
                ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_before = Pt(6); ip.paragraph_format.space_after = Pt(3)
                ip.paragraph_format.keep_with_next = True
                w, h = img_w_h(path_, CONTENT_W_CM * 0.86)
                ip.add_run().add_picture(path_, width=w, height=h)
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cp.paragraph_format.space_after = Pt(9)
                r = cp.add_run(f"{L_FIG} {lab}. "); r.bold = True; r.font.size = Pt(8.5)
                r.font.name = "Times New Roman"
                add_rich(cp, cap, 8.5)

        elif k == "table":
            cap, lab, heads, rows, note = blk[1], blk[2], blk[3], blk[4], blk[5]
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cp.paragraph_format.space_before = Pt(6); cp.paragraph_format.space_after = Pt(3)
            cp.paragraph_format.keep_with_next = True
            r = cp.add_run(f"{L_TAB} {lab}. "); r.bold = True; r.font.size = Pt(8.5)
            r.font.name = "Times New Roman"
            add_rich(cp, cap, 8.5)

            tb = doc.add_table(rows=1, cols=len(heads))
            tb.alignment = WD_TABLE_ALIGNMENT.LEFT
            for j, htxt in enumerate(heads):
                c = tb.cell(0, j); c.paragraphs[0].paragraph_format.space_after = Pt(1)
                add_rich(c.paragraphs[0], htxt, 8.2)
                for run in c.paragraphs[0].runs:
                    run.bold = True
                cell_border(c, "top", 10, "333333"); cell_border(c, "bottom", 6, "333333")
            for i, row in enumerate(rows):
                cells = tb.add_row().cells
                for j, val in enumerate(row):
                    cells[j].paragraphs[0].paragraph_format.space_after = Pt(1)
                    add_rich(cells[j].paragraphs[0], str(val), 8.2)
                    if i % 2 == 1:
                        shade(cells[j], "FAFAFA")
                    if i == len(rows) - 1:
                        cell_border(cells[j], "bottom", 10, "333333")
            if note:
                np_ = para(doc, note, size=7.9, align=WD_ALIGN_PARAGRAPH.LEFT,
                           space_after=10, color=GREY)
                np_.paragraph_format.space_before = Pt(3)

        elif k == "refs":
            for ref in blk[1]:
                rp = doc.add_paragraph(); rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf = rp.paragraph_format
                pf.left_indent = Cm(0.7); pf.first_line_indent = Cm(-0.7)
                pf.space_after = Pt(3.6); pf.line_spacing = 1.05
                add_rich(rp, ref, 8.7)

    doc.save(path)
    return path


from build_pdf import STEM

if __name__ == "__main__":
    p = build(os.path.join(OUT, f"{STEM.get(LANG, 'paper-' + LANG)}.docx"))
    print(f"DOCX [{LANG}] →", p, f"({os.path.getsize(p)/1024:.0f} KB)")
